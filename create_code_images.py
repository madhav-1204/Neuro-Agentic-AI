from PIL import Image, ImageDraw, ImageFont
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import RawTokenFormatter
from pygments.token import Token
import os
from docx import Document
from docx.shared import Inches

# Read code files
files = [
    "b:/testp/backend/main.py",
    "b:/testp/backend/app/core/vision_agent.py",
    "b:/testp/backend/app/core/reasoning_agent.py",
    "b:/testp/backend/app/core/gemini_vision_agent.py",
    "b:/testp/backend/app/core/explainability_agent.py",
]

# Color scheme (VS Code dark)
COLORS = {
    'bg': (30, 30, 30),
    'text': (212, 212, 212),
    'keyword': (86, 156, 214),
    'string': (206, 145, 120),
    'comment': (87, 166, 74),
    'number': (181, 206, 168),
    'function': (220, 220, 170),
}

def get_color(token_type):
    if token_type in Token.Keyword:
        return COLORS['keyword']
    elif token_type in Token.String:
        return COLORS['string']
    elif token_type in Token.Comment:
        return COLORS['comment']
    elif token_type in Token.Number:
        return COLORS['number']
    elif token_type in Token.Name.Function:
        return COLORS['function']
    return COLORS['text']

def code_to_image(code_text, filename):
    """Convert code to styled image"""
    # Try to use monospace font
    try:
        font = ImageFont.truetype("C:\\Windows\\Fonts\\consola.ttf", 11)
        line_height = 20
    except:
        font = ImageFont.load_default()
        line_height = 10

    lines = code_text.split('\n')

    # Calculate dimensions
    max_width = max([len(line) * 8 for line in lines[:50]]) + 40  # Limit preview
    height = min(len(lines[:50]) * line_height + 40, 1200)

    # Create image
    img = Image.new('RGB', (min(max_width, 1000), height), COLORS['bg'])
    draw = ImageDraw.Draw(img)

    y = 20
    for i, line in enumerate(lines[:50]):  # Limit to 50 lines
        # Simple tokenization for visual effect
        tokens = highlight(line, PythonLexer(), RawTokenFormatter())
        lexer = PythonLexer()
        tokens = list(lexer.get_tokens(line))

        x = 20
        for token_type, value in tokens:
            color = get_color(token_type)
            draw.text((x, y), value, fill=color, font=font)
            x += len(value) * 8

        y += line_height
        if y > height - 40:
            break

    # Add filename watermark
    draw.text((20, height - 25), f"← {filename}", fill=(100, 100, 100), font=font)

    return img

# Create Word document
doc = Document()

for file_path in files:
    filename = os.path.basename(file_path)

    # Read code
    with open(file_path, 'r', encoding='utf-8') as f:
        code = f.read()

    # Add title
    doc.add_heading(filename, level=1)

    # Generate and add image
    img = code_to_image(code, filename)
    img_path = f"b:/testp/temp_{filename}.png"
    img.save(img_path)

    # Add image to document
    doc.add_picture(img_path, width=Inches(6))

    # Clean up temp image
    os.remove(img_path)

    # Add page break
    doc.add_page_break()

# Save document
output = "b:/testp/Code_Snapshots.docx"
doc.save(output)
print(f"Saved to {output}")

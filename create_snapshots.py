from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Define files and their content
files_data = [
    ("main.py", """import os
import logging
from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse
from api.health import router as health_router
from api.analyze import router as analyze_router
from fastapi.middleware.cors import CORSMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from app.database import init_db

# ── Logging ─────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s │ %(levelname)-8s │ %(name)s │ %(message)s",
)
logger = logging.getLogger(__name__)

# ── Rate limiter ────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address, default_limits=["60/minute"])

# ── App ─────────────────────────────────────────────────────────────
app = FastAPI(
    title="Neuro Diagnosis Backend",
    version="2.0.0",
    description="AI-powered brain MRI analysis API",
)

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ── CORS ────────────────────────────────────────────────────────────
_extra_origins = os.getenv("CORS_ORIGINS", "")
allowed_origins = ["http://localhost:5173", "http://localhost:5174"]
if _extra_origins:
    allowed_origins.extend([o.strip() for o in _extra_origins.split(",") if o.strip()])

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Global exception handler ───────────────────────────────────────
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error("Unhandled error on %s %s: %s", request.method, request.url, exc)
    return JSONResponse(status_code=500, content={"error": "Internal server error"})

# ── Request logging middleware ──────────────────────────────────────
@app.middleware("http")
async def log_requests(request: Request, call_next):
    logger.info("%s %s", request.method, request.url.path)
    response = await call_next(request)
    return response

# ── Startup ─────────────────────────────────────────────────────────
@app.on_event("startup")
def on_startup():
    init_db()
    logger.info("Database initialised, server ready.")

# ── Routers ─────────────────────────────────────────────────────────
app.include_router(health_router)
app.include_router(analyze_router)

@app.get("/")
def root():
    return {"message": "Backend is running", "version": "2.0.0"}
"""),

    ("vision_agent.py", """import torch
import torch.nn as nn
from torchvision import models, transforms
from PIL import Image, UnidentifiedImageError

from app.config.settings import (
    MODEL_ARCHITECTURE,
    MODEL_PATHS,
    NUM_CLASSES,
    CLASS_NAMES,
    IMAGE_SIZE,
    MEAN,
    STD,
)

class VisionAgent:
    def __init__(self):
        self.model = None
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.transform = self._get_transform()

    def _get_transform(self):
        return transforms.Compose([
            transforms.Resize(IMAGE_SIZE),
            transforms.ToTensor(),
            transforms.Normalize(mean=MEAN, std=STD),
        ])

    def _build_model(self):
        if MODEL_ARCHITECTURE == "resnet18":
            model = models.resnet18(weights=None)
            model.fc = nn.Linear(model.fc.in_features, NUM_CLASSES)
        elif MODEL_ARCHITECTURE == "efficientnet_b3":
            model = models.efficientnet_b3(weights=None)
            model.classifier[1] = nn.Linear(model.classifier[1].in_features, NUM_CLASSES)
        else:
            raise ValueError(f"Unsupported model: {MODEL_ARCHITECTURE}")
        return model

    def load_model(self):
        if self.model is not None:
            return self.model
        model_path = MODEL_PATHS[MODEL_ARCHITECTURE]
        self.model = self._build_model()
        checkpoint = torch.load(model_path, map_location=self.device)
        self.model.load_state_dict(checkpoint)
        self.model.to(self.device)
        self.model.eval()
        return self.model

    def preprocess_image(self, image_path):
        try:
            image = Image.open(image_path).convert("RGB")
        except UnidentifiedImageError:
            raise ValueError(f"File is not a valid image: {image_path}")
        tensor = self.transform(image).unsqueeze(0)
        return tensor.to(self.device), image

    def predict(self, image_path):
        if self.model is None:
            self.load_model()
        tensor, original_image = self.preprocess_image(image_path)
        with torch.no_grad():
            logits = self.model(tensor)
            probabilities = torch.softmax(logits, dim=1).cpu().numpy()[0]
        predicted_idx = int(probabilities.argmax())
        predicted_class = CLASS_NAMES[predicted_idx]
        confidence = float(probabilities[predicted_idx] * 100)
        return {
            "probabilities": probabilities,
            "predicted_class": predicted_class,
            "predicted_idx": predicted_idx,
            "confidence": confidence,
            "original_image": original_image,
            "class_names": CLASS_NAMES,
            "preprocessed_tensor": tensor,
        }

    def get_model_for_gradcam(self):
        if self.model is None:
            self.load_model()
        return self.model
"""),

    ("reasoning_agent.py", """import os
import logging
from app.config.settings import GEMINI_MODEL

logger = logging.getLogger(__name__)

try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

REASONING_PROMPT = (
    "You are a neuroradiology AI assistant. Given the following brain MRI "
    "classification result, provide a concise 3-4 sentence medical explanation "
    "covering: (1) what the predicted condition means, (2) its clinical "
    "significance, and (3) recommended next steps. Be professional and factual."
)

class ReasoningAgent:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        self.client = None
        self.model = None
        if self.api_key and GEMINI_AVAILABLE:
            self.client = genai.Client(api_key=self.api_key)
            self.model = GEMINI_MODEL

    def generate_explanation(self, predicted_class, confidence, probabilities, class_names, image_path=None):
        if not self.client:
            return self._fallback(predicted_class, confidence)
        try:
            prob_summary = ", ".join(
                f"{cls}: {prob*100:.1f}%" for cls, prob in zip(class_names, probabilities)
            )
            prompt = REASONING_PROMPT.format(
                predicted_class=predicted_class,
                confidence=confidence,
                prob_summary=prob_summary,
            )
            response = self.client.models.generate_content(
                model=self.model,
                contents=[types.Content(parts=[types.Part.from_text(text=prompt)])],
            )
            explanation = response.text.strip()
            return explanation if explanation else self._fallback(predicted_class, confidence)
        except Exception as e:
            logger.error("ReasoningAgent Gemini call failed: %s", e)
            return self._fallback(predicted_class, confidence)

    def _fallback(self, predicted_class, confidence):
        return (f"The model predicts {predicted_class} with {confidence:.1f}% confidence. "
                f"Please consult a qualified neuroradiologist for a definitive diagnosis.")
"""),

    ("gemini_vision_agent.py", """import os
import json
import re
import time
import logging
from app.config.settings import GEMINI_MODEL

logger = logging.getLogger(__name__)

try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

class GeminiVisionAgent:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        self.client = None
        self.model = GEMINI_MODEL
        if self.api_key and GEMINI_AVAILABLE:
            self.client = genai.Client(api_key=self.api_key)

    @staticmethod
    def _read_image_bytes(image_path):
        ext = os.path.splitext(image_path)[1].lower()
        media_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
        mime_type = media_map.get(ext, "image/jpeg")
        with open(image_path, "rb") as f:
            data = f.read()
        return data, mime_type

    def analyze(self, image_path):
        if not self.client:
            raise RuntimeError("Gemini client not initialised. Set GEMINI_API_KEY in .env file.")
        image_bytes, mime_type = self._read_image_bytes(image_path)
        models_to_try = [self.model, "gemini-2.5-flash-lite"]
        last_error = None

        for model_name in models_to_try:
            for attempt in range(4):
                try:
                    logger.info(f"Gemini request: model={model_name}, attempt={attempt+1}")
                    response = self.client.models.generate_content(model=model_name, contents=content)
                    raw = response.text.strip()
                    raw = re.sub(r"^```(?:json)?\\\\s*", "", raw)
                    raw = re.sub(r"\\\\s*```$", "", raw)
                    try:
                        return json.loads(raw)
                    except json.JSONDecodeError:
                        return {"error": "Failed to parse Gemini response", "raw": raw}
                except Exception as e:
                    last_error = e
                    if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                        wait = (2 ** attempt) * 2
                        logger.warning(f"Rate limited. Waiting {wait}s...")
                        time.sleep(wait)
                        continue
                    else:
                        raise

        raise RuntimeError(f"All models exhausted. Last error: {last_error}")
"""),

    ("explainability_agent.py", """import torch
import numpy as np
from pytorch_grad_cam import GradCAM
from pytorch_grad_cam.utils.image import show_cam_on_image
from pytorch_grad_cam.utils.model_targets import ClassifierOutputTarget
from app.config.settings import IMAGE_SIZE

class ExplainabilityAgent:
    def __init__(self, model):
        self.model = model

    def _get_target_layer(self):
        return [self.model.layer4[-1]]

    def generate_gradcam(self, image_tensor, predicted_idx):
        with GradCAM(model=self.model, target_layers=self._get_target_layer()) as cam:
            targets = [ClassifierOutputTarget(predicted_idx)]
            grayscale_cam = cam(input_tensor=image_tensor, targets=targets)[0]
        return grayscale_cam

    def create_overlay(self, original_image, grayscale_cam):
        img = original_image.resize(IMAGE_SIZE)
        rgb_img = np.array(img).astype(np.float32) / 255.0
        return show_cam_on_image(rgb_img, grayscale_cam, use_rgb=True)

    def explain(self, image_tensor, original_image, predicted_idx):
        cam = self.generate_gradcam(image_tensor, predicted_idx)
        overlay = self.create_overlay(original_image, cam)
        return {
            "heatmap": cam,
            "overlay": overlay,
            "explanation": "Highlighted regions influenced the prediction.",
        }
"""),
]

# Create document with one file per page
for idx, (filename, code) in enumerate(files_data):
    # Add title
    title = doc.add_heading(filename, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add code
    code_paragraph = doc.add_paragraph()
    code_run = code_paragraph.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)

    # Add page break except after last file
    if idx < len(files_data) - 1:
        doc.add_page_break()

# Save document
output_path = 'b:/testp/Code_Snapshots.docx'
doc.save(output_path)
print(f"Document saved: {output_path}")
